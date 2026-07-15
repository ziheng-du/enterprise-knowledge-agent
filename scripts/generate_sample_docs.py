"""生成多格式演示文档（PDF / DOCX），供入库与 README 演示。

用法（项目根目录）：
    python scripts/generate_sample_docs.py
"""

from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document as DocxDocument

RAW_DIR = Path(__file__).resolve().parent.parent / "data" / "raw_docs"


def write_docx() -> Path:
    """写入《IT设备领用须知.docx》（public 密级）。"""
    path = RAW_DIR / "IT设备领用须知.docx"
    doc = DocxDocument()
    doc.add_heading("IT 设备领用须知", level=1)
    doc.add_paragraph(
        "本文档为演示用 Word（.docx）样例，用于验证 document_loader 与入库链路对 Office 格式的支持。"
    )
    doc.add_heading("领用范围", level=2)
    doc.add_paragraph(
        "新员工入职后可申请笔记本电脑一台、显示器一台（按岗位标准配置）。"
        "外设（键鼠、耳机）按需申领，单人累计外设领用金额不超过 800 元/年。"
    )
    doc.add_heading("归还与离职", level=2)
    doc.add_paragraph(
        "员工离职或调岗时须在最后一个工作日前归还全部公司资产，"
        "由 IT 资产管理员验收并在资产系统中注销。"
        "损坏或遗失按公司固定资产管理办法折价赔偿。"
    )
    doc.add_heading("咨询渠道", level=2)
    doc.add_paragraph("IT 服务台邮箱：it-helpdesk@example.com（模拟地址）。")
    doc.save(str(path))
    return path


def write_pdf() -> Path:
    """写入可抽取文本的简易 PDF《消防与安全须知.pdf》（public）。

    使用 Type1/Helvetica，正文为英文条款（证明 PDF 入库链路）；
    中文制度样例见同目录 md/docx。扫描件/OCR 不在本项目范围。
    """
    path = RAW_DIR / "消防与安全须知.pdf"
    lines = [
        "Fire Safety Notice - Sample PDF for ingest demo",
        "1. Evacuation drills: twice per year.",
        "2. Fire extinguisher inspection: every 90 days.",
        "3. Emergency exits must stay clear at all times.",
        "4. Report hazards to EHS within 24 hours.",
        "Keywords: fire safety, extinguisher, evacuation, EHS.",
    ]
    content = "BT /F1 11 Tf 50 750 Td 16 TL\n"
    for i, line in enumerate(lines):
        safe = line.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
        content += f"({safe}) Tj\n" if i == 0 else f"T* ({safe}) Tj\n"
    content += "ET\n"
    stream = content.encode("latin-1")

    objects: list[bytes] = [
        b"1 0 obj<< /Type /Catalog /Pages 2 0 R >>endobj\n",
        b"2 0 obj<< /Type /Pages /Kids [3 0 R] /Count 1 >>endobj\n",
        (
            b"3 0 obj<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>endobj\n"
        ),
        f"4 0 obj<< /Length {len(stream)} >>stream\n".encode("ascii")
        + stream
        + b"\nendstream\nendobj\n",
        b"5 0 obj<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>endobj\n",
    ]

    pdf = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for obj in objects:
        offsets.append(len(pdf))
        pdf.extend(obj)
    xref_pos = len(pdf)
    pdf.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    pdf.extend(b"0000000000 65535 f \n")
    for off in offsets[1:]:
        pdf.extend(f"{off:010d} 00000 n \n".encode("ascii"))
    pdf.extend(
        (
            f"trailer<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_pos}\n%%EOF\n"
        ).encode("ascii")
    )
    path.write_bytes(pdf)
    return path


def main() -> None:
    """生成样例并打印路径。"""
    RAW_DIR.mkdir(parents=True, exist_ok=True)
    docx_path = write_docx()
    pdf_path = write_pdf()
    print(f"已生成: {docx_path}")
    print(f"已生成: {pdf_path}")


if __name__ == "__main__":
    main()
