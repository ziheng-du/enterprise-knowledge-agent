"""文档加载模块：将 raw_docs 目录下的企业制度文档加载为 LangChain Document。

支持格式：.md / .txt / .pdf / .docx。
每个 Document 的 metadata 统一包含：
- source: 来源文件名（不含路径），用于回答中标注引用来源
- 位置信息: pdf 为 page（页码，从 0 起），docx/txt/md 为整文档（切分阶段补充块序号）

设计说明：加载逻辑独立于切分与检索，单个文件解析失败只记录日志并跳过，
不中断整批入库流程。
"""

from pathlib import Path

from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader, TextLoader

from app.rag.access_control import resolve_access_level
from app.utils.logger import get_logger

logger = get_logger(__name__)

# 支持的文件扩展名（小写）
SUPPORTED_EXTENSIONS = {".md", ".txt", ".pdf", ".docx"}


def _load_text_file(file_path: Path) -> list[Document]:
    """加载 .md / .txt 纯文本文件（utf-8 编码）。"""
    return TextLoader(str(file_path), encoding="utf-8").load()


def _load_pdf_file(file_path: Path) -> list[Document]:
    """加载 .pdf 文件，每页一个 Document，metadata 自带 page 页码。"""
    return PyPDFLoader(str(file_path)).load()


def _load_docx_file(file_path: Path) -> list[Document]:
    """加载 .docx 文件。

    说明：langchain 的 Docx2txtLoader 依赖 docx2txt 包（不在本项目依赖中），
    这里直接用已有依赖 python-docx 做轻量解析，按段落拼接为单个 Document。
    """
    # 延迟导入：仅在实际加载 docx 时才需要该依赖
    from docx import Document as DocxDocument

    docx = DocxDocument(str(file_path))
    # 过滤空段落，保留正文结构（段落间以换行分隔，便于后续按段落边界切分）
    text = "\n".join(p.text for p in docx.paragraphs if p.text.strip())
    return [Document(page_content=text, metadata={"source": str(file_path)})]


# 扩展名 -> 加载函数 的分派表
_LOADER_DISPATCH = {
    ".md": _load_text_file,
    ".txt": _load_text_file,
    ".pdf": _load_pdf_file,
    ".docx": _load_docx_file,
}


def load_single_document(file_path: Path) -> list[Document]:
    """加载单个文档文件并规范化元数据。

    Args:
        file_path: 文档文件路径。

    Returns:
        Document 列表（pdf 每页一个，其余格式整文件一个）。
        metadata.source 统一规范化为文件名（不含目录路径）。

    Raises:
        ValueError: 文件扩展名不受支持时抛出。
    """
    ext = file_path.suffix.lower()
    if ext not in _LOADER_DISPATCH:
        raise ValueError(f"不支持的文件格式: {ext}（支持 {sorted(SUPPORTED_EXTENSIONS)}）")

    docs = _LOADER_DISPATCH[ext](file_path)
    # 统一 source 为文件名：加载器默认写入完整路径，展示引用来源时只需文件名
    access_level = resolve_access_level(file_path.name)
    for doc in docs:
        doc.metadata["source"] = file_path.name
        # 密级写入 metadata，切分后随块保留，供检索阶段按角色过滤
        doc.metadata["access_level"] = access_level
    return docs


def load_documents(docs_dir: Path) -> list[Document]:
    """批量加载目录下所有受支持的文档。

    Args:
        docs_dir: 原始文档目录（通常为 settings.raw_docs_dir）。

    Returns:
        所有成功加载的 Document 列表。

    边界处理：
    - 目录不存在或为空：记录 warning，返回空列表，由调用方决定后续行为
    - 不支持的扩展名：跳过并记录 debug 日志
    - 单个文件解析失败：记录 error 后跳过，不中断整批加载
    """
    if not docs_dir.exists():
        logger.warning("文档目录不存在: %s", docs_dir)
        return []

    documents: list[Document] = []
    for file_path in sorted(docs_dir.iterdir()):
        if not file_path.is_file():
            continue
        if file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            logger.debug("跳过不支持的文件: %s", file_path.name)
            continue
        try:
            docs = load_single_document(file_path)
            documents.extend(docs)
            logger.info("已加载 %s（%d 个文档片段）", file_path.name, len(docs))
        except Exception:
            # 单文件失败不影响整批入库，异常堆栈记入日志便于排查
            logger.exception("加载文件失败，已跳过: %s", file_path.name)

    if not documents:
        logger.warning("目录 %s 中没有加载到任何文档", docs_dir)
    return documents
